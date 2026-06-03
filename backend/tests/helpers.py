"""Shared test document builders for parser tests and perf scripts."""
from __future__ import annotations

import io

from docx import Document


def build_sample_docx(question_count: int = 10) -> bytes:
    doc = Document()
    for i in range(1, question_count + 1):
        doc.add_paragraph("[单选]")
        doc.add_paragraph(f"{i}. 性能测试题 {i}：2+2 等于几？")
        doc.add_paragraph("A. 3")
        doc.add_paragraph("B. 4")
        doc.add_paragraph("C. 5")
        doc.add_paragraph("答案: B")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def write_fixture_files() -> dict[str, bytes]:
    """Return named payloads used by Locust and baseline scripts."""
    return {
        "small_10q.docx": build_sample_docx(10),
        "medium_50q.docx": build_sample_docx(50),
    }
